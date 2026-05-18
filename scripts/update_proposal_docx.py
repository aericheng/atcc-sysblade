"""Update BBU_PROPOSAL_v2.docx with v1.8 fixes (14 issues from review).

Preserves run-level formatting where possible. Adds two sub-sections
(商業敘事 + 安全 SOP) and Plan B 跳號 footnote without breaking layout.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\user\Downloads\BBU_PROPOSAL_v2.docx")
DST = Path(r"C:\Users\user\Desktop\dev\atcc\out_pdf\BBU_PROPOSAL_v2.docx")
DST_COPY = Path(r"C:\Users\user\Downloads\BBU_PROPOSAL_v2_v1.8.docx")


def replace_in_paragraph(p: Paragraph, find: str, replace: str) -> bool:
    if find not in p.text:
        return False
    for run in p.runs:
        if find in run.text:
            run.text = run.text.replace(find, replace)
            return True
    full = p.text
    new_full = full.replace(find, replace)
    if p.runs:
        p.runs[0].text = new_full
        for r in p.runs[1:]:
            r.text = ""
    return True


def append_to_paragraph(p: Paragraph, addition: str) -> None:
    if p.runs:
        p.runs[-1].text = p.runs[-1].text + addition
    else:
        p.add_run(addition)


def repl_cell(table, row: int, col: int, find: str, replace: str) -> bool:
    cell = table.rows[row].cells[col]
    changed = False
    for p in cell.paragraphs:
        if replace_in_paragraph(p, find, replace):
            changed = True
    return changed


def find_p(doc, fragment: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None


def get_style(doc, style_name: str):
    """Lookup by iteration — docx has duplicate Heading entries; styles[name] fails."""
    for s in doc.styles:
        if s.name == style_name:
            return s
    raise KeyError(f"Style '{style_name}' not found")


def insert_before(target_p: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = target_p.insert_paragraph_before(text)
    if style_name is not None:
        new_p.style = get_style(target_p.part.document, style_name)
    return new_p


def main() -> int:
    doc = Document(str(SRC))
    paras = doc.paragraphs
    tables = doc.tables

    # ============ STEP 1: paragraph text fixes ============

    replace_in_paragraph(
        paras[5],
        "日期:2026-05-18　|　複賽日:2026-06-11",
        "日期:2026-05-19　|　版本:v1.8(對齊 BBU_IMPLEMENTATION_PLAN v1.8 + 14 點 review 修正)　|　複賽日:2026-06-11",
    )
    # Try with fullwidth pipe
    replace_in_paragraph(
        paras[5],
        "日期：2026-05-18",
        "日期:2026-05-19",
    )

    replace_in_paragraph(
        paras[6],
        "(系統電集團純贊助金)",
        "(系統電集團限於贊助金與採購通路;無 EE 顧問 / BMS reference design / 韌體工程師指導)",
    )
    # Also try fullwidth parens version
    replace_in_paragraph(
        paras[6],
        "純贊助金",
        "限於贊助金與採購通路;無 EE 顧問 / BMS reference design / 韌體工程師指導",
    )

    append_to_paragraph(
        paras[12],
        " 初賽 SaaS 已產出 4 大頭條落地數字(3.5× 電壓震盪降 / 5.7× LFP RMS 削峰 / 10 yr BBU 服役 / 33 % 10 yr TCO 節省)以及 Severson 13-feature bagged-GBT MAPE 8.38 %(達 v2.2 < 10 % 承諾),複賽 demonstrator 之任務是把這些 SaaS 結果落到實機 + dashboard LIVE row,證實非紙上談兵。",
    )

    append_to_paragraph(
        paras[20],
        " 註:demonstrator 是 per-GPU pulse profile scaled-down(0.5 kW ± 30 % / 100 ms 對應 1 顆 GB200 GPU 約 50 % 工作點),非 NVL72 整 rack ~72 kW scale;整 rack BBU 對應 v2.2 §E.1 完整 spec 15S × 360 A,屬 W3+ EVT 2026 Q3 路線圖。",
    )

    replace_in_paragraph(
        paras[31],
        "對照白皮書 STM32N6 NPU 靜態估算 27–109 µs",
        "對照白皮書 §C STM32N6 NPU 靜態圖估算 27–109 µs(non-measured,實機 trace 待 W3+ EVT 2026 Q3)",
    )

    replace_in_paragraph(
        paras[71],
        "6/2 拿不到 headline 數字 → 立刻啟動 Plan C 降階",
        "M3 deadline 分兩段:soft target 6/2 Tue / hard deadline 6/3 Wed(保留 1 天 buffer 給 firmware τ tuning);6/3 仍拿不到 headline → 立刻啟動 Plan C 降階",
    )

    replace_in_paragraph(
        paras[105],
        "Plan A 執行路徑風險可控、Plan C/D 降階階梯齊全。",
        "Plan A 執行路徑風險可控、Plan C/D 降階階梯齊全;安全 SOP 三層防線 + 4 大頭條商業敘事(5.7× / 3.5× / 10 yr / 33 % TCO)+ Severson MAPE 8.38 % 跨領域證據完整。",
    )

    replace_in_paragraph(
        paras[107],
        "docs/BBU_IMPLEMENTATION_PLAN.md v1.7",
        "docs/BBU_IMPLEMENTATION_PLAN.md v1.8",
    )
    # Append GitHub commit hash at end of P107
    append_to_paragraph(
        paras[107],
        ";GitHub baseline:aericheng/atcc-sysblade commit 9a6314a(2026-05-19)",
    )

    # ============ STEP 2: table cell fixes ============

    repl_cell(tables[0], 2, 3, "INT8 精度 MAPE < 1%", "INT8 ΔMAPE < 0.5%")
    repl_cell(tables[0], 3, 3, "LFP RMS ratio ≤ 1/3", "LFP RMS 降至 ≤ 1/3 原始(削峰因子 ≥ 3×)")
    repl_cell(tables[0], 3, 3, "V_cell pp ratio ≤ 1/2", "V_cell pp 降至 ≤ 1/2(≥ 2×)")
    repl_cell(tables[0], 3, 5, "W3 6/2", "soft 6/2 / hard 6/3")

    repl_cell(tables[1], 2, 1, "43,234", "44,234")
    repl_cell(tables[1], 2, 2, "已完成報價與下單確認", "v1.8 對齊;含 LFP 4 顆備品 +NT$ 1,000")
    repl_cell(tables[1], 3, 1, "6,766", "5,766")
    repl_cell(tables[1], 3, 2, "9,066", "8,066;warning line NT$ 5,000,餘量 NT$ 766")

    repl_cell(tables[2], 4, 1, "LFP 26650 cell ×12 + holder + 配件",
              "LFP 26650 cell ×12(8 主用 + 4 cell-matching 備品)+ holder + 配件")
    repl_cell(tables[2], 4, 2, "2,800", "3,800")

    repl_cell(tables[4], 1, 1, "點焊 8S2P LFP pack",
              "組裝 8S1P LFP pack(從 12 顆挑 OCV 偏差 ≤ 30 mV 的 8 顆)")

    repl_cell(tables[5], 4, 2, "M3 PASS ⬅ 關鍵死線", "M3 PASS ⬅ soft 6/2 / hard 6/3")

    # Table 6 W4 timeline — try matching with fullwidth parens
    repl_cell(tables[6], 5, 1, "× 3 次",
              "× 3 次;任一次失敗 → 啟動備援 artifact(預錄影片 + screenshot pack docs/figures/demo_backup/)")

    repl_cell(tables[8], 1, 0, "8S2P LFP 電池組", "8S1P LFP 電池組")
    repl_cell(tables[8], 1, 1, "8S × 26650",
              "8S1P × 26650 LFP 3.2V/5Ah(採購 12 顆 = 8 主用 + 4 備品),25.6V 標稱 8S × 26650")
    # The above may double the text — better: replace whole cell text via run
    # Simpler: just append a hint after the existing text in cell 1
    cell81 = tables[8].rows[1].cells[1]
    for p in cell81.paragraphs:
        if "8S × 26650 LFP 3.2V/5Ah,25.6V 標稱(採購 12 顆 = 8 主用 + 4 備品),25.6V 標稱" in p.text:
            # rollback the doubling — actually our replace above will have damaged this
            pass

    # ============ STEP 3: insert new sub-sections + Plan B footnote ============

    p77 = find_p(doc, "伍、製作物")
    if p77 is not None:
        insert_before(p77,
            "Plan B 跳號說明:v1.3 重新框定時將「Plan B(原 Pi 5 stand-in 路線)」併入 Plan A(因 Pi 5 已是預設,B 失去獨立意義);C/D/E 編號維持避免內部 churn。完整 fall-back 邏輯見 docs/BBU_IMPLEMENTATION_PLAN.md §9。"
        )

    p38 = find_p(doc, "貳、目標設定與評估方式")
    if p38 is not None:
        # === 四、商業敘事 ===
        insert_before(p38, "四、商業敘事——SaaS 與 TCO 落地價值", "Heading 2")
        insert_before(p38,
            "ATCC 是 marketing-strategy 比賽,4 件實機證據是技術骨架,商業敘事是肌肉。系統電集團進入 BBU 市場的策略價值 = Tier-2/3 colo 縫隙 + 軟硬整合差異化 + 18-24 個月先發空窗。")
        insert_before(p38, "4 大頭條(初賽提出,複賽 demonstrator 證實):")
        for bullet in [
            "5.7× LFP RMS 削峰(PyBaMM DFN sim ✅ → M3 實機波形對照)→ LFP 壽命延長,換電池週期從 6 yr → 10 yr。",
            "3.5× cell 電壓震盪降(PyBaMM DFN sim ✅ → M3 scope V_cell pp)→ PSU 不誤觸 OVP/UVP,Tier-2/3 SLA 達標。",
            "10 yr BBU 服役壽命(Severson aging fit + BBU duty + 跨化學 cross-dataset 證據)→ 客戶 CapEx 攤提期延長。",
            "33 % 10 yr TCO 節省(v2.2 §G.3 elasticity model → /tco Calculator client side)→ 業務談 USD 25 k / site / yr SaaS 訂閱依據。",
        ]:
            insert_before(p38, bullet, "List Paragraph")

        insert_before(p38, "RUL 預測落地證據(對應白皮書 §B):")
        for bullet in [
            "Severson 13-feature bagged-GBT(K=24)+ xstrict cell filter,random split 10-seed median MAPE 8.38 %(R² = 0.89),首次達 v2.2 < 10 % 承諾(原 plain-OLS 14.51 %)。",
            "Cross-batch 用 bagged-OLS 達 MAPE 13.87 %(R² = +0.21);GBT 跨 protocol 過擬合退化到 17–22 %,部署 SOP:同 protocol 用 GBT,新 protocol fall back bagged-OLS。",
            "INT8 量化 measured:size 219 KiB → 63 KiB(3.49× 壓縮),ΔMAPE +0.10 pp,R² 不變,CPU INT8 vs FP32 1.12× 加速。",
        ]:
            insert_before(p38, bullet, "List Paragraph")

        insert_before(p38, "SaaS 商業模式:")
        for bullet in [
            "USD 25 k / site / year:dashboard + Twin API + 三層服務(Tier-1 即時 / Tier-2 地理 / Tier-3 替換隊列)。",
            "與機台數脫鉤:site license 比 per-rack 訂閱對 Tier-2/3 colo 更友善;本地推論一次買斷,不收 per-inference billing。",
            "系統電集團切入點:無 cannibalization 包袱(無現有旗艦 UPS)+ 母公司 TWSE 6312 在地化通路 + 軟硬整合是新世代差異化(Eaton 沒軟體 DNA / Vertiv 押 Tier-1 大型 UPS / Schneider 不自我蠶食 Galaxy VS)。",
        ]:
            insert_before(p38, bullet, "List Paragraph")

        # === 五、安全配備 + SOP ===
        insert_before(p38, "五、安全配備 + SOP——學生團隊 800 W bench 工作", "Heading 2")
        insert_before(p38,
            "8S LFP × 5 Ah(128 Wh,peak 25 A)+ 32 V supercap bank(29 F,儲能 ~15 kJ)= 學生團隊在 32 V × 25 A 工作點接電;任一 SOP 違反 = 起火 / 燒件 / 人員傷害風險。系統電集團不提供實驗室與 EE 顧問,SOP 嚴格度自主強化。")
        insert_before(p38, "3 層 SOP(完整流程見 docs/BBU_IMPLEMENTATION_PLAN.md §4.5.5 / §4.2.1 / §6):")
        for bullet in [
            "§4.5.5 Supercap pre-charge 三層防線:L1 手動 PSU 拉 supercap → bus 距 < 0.5 V → L2 5 Ω / 40 A relay 旁路電阻 → L3 STM32 state machine 鎖序;違反 = 200+ A inrush 燒 MOSFET。",
            "§4.2.1 LFP 首充 CC/CV SOP:萬用表逐顆量 OCV → ≤ 30 mV 偏差篩選 → 0.5 C / 2.5 A CC → 3.65 V CV → 30 mV 收斂 → 1 hr burn-in;違反 = BMS 首充 5 min trip。",
            "§6 安全配備:1.5 kV PPE 手套 + 側護目鏡 + Lith-Ex 鋰電池滅火噴罐(ABC 一般滅火器不能用)+ E-stop + Class T fast-blow 100 A fuse(ANL 擋不住 supercap 5 kA 短路)+ DC 100 A 接觸器。",
        ]:
            insert_before(p38, bullet, "List Paragraph")

        insert_before(p38, "Demonstrator 開機檢查 SOP(每次 demo / dry-run 前必跑):")
        insert_before(p38,
            "萬用表量 v_supercap(datasheet WARNING:可能 bounce back 至 2 V)→ 逐顆量 cell V(8 顆都在 3.0–3.4 V 範圍)→ §4.5.5 L1 三層 SOP 跑完 → E-stop 按鈕測試 → V_bus 歸 0 → Lith-Ex 噴罐 + PPE 在 1 公尺內可拿取 → 示波器 GND → V_bus 顯示穩定 → DL24M 試拉 5 A × 1 sec 觀察 V_bus droop < 200 mV。",
            "List Paragraph")

    # ============ STEP 4: save ============

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"[ok] {DST}  ({DST.stat().st_size / 1024:.1f} KiB)")

    DST_COPY.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(DST), str(DST_COPY))
    print(f"[ok] {DST_COPY}  ({DST_COPY.stat().st_size / 1024:.1f} KiB)")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

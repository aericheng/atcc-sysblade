"""Generate Sysblade_HyperBuffer_Proposal_v2.2.docx by surgical edits on v2.1.docx.

Why a surgical-edit script (vs. python-docx default styles): v2.1.docx was
authored in Microsoft Word with custom heading colours, paragraph spacing,
table cell borders. python-docx defaults won't match. We load v2.1.docx as
a template, mutate specific paragraphs / cells, and save — every untouched
style is preserved.

Inputs
------
- v2.1.docx at NYCU/ATCC/ (the authoritative source authored by the team's
  business member). This script only needs READ access; the file is not
  modified in place.

Outputs (both written)
----------------------
- docs/proposal_v2.2_additions/Sysblade_HyperBuffer_Proposal_v2.2.docx
- C:/Users/user/Desktop/NYCU/ATCC/Sysblade_HyperBuffer_Proposal_v2.2.docx

What v2.2 adds vs v2.1
----------------------
1. Cover: v2.2 修訂版 (2026-05-06) + Live demo + GitHub repo URL block.
2. Abstract: insert measured-numbers callout (8.38 % MAPE, +0.10 pp INT8,
   η = 0.945 rainflow, drilldown shipped, repo public).
3. New 附件 D. 技術交付物實證 (D.1–D.6 with 6 tables / paragraphs covering
   ML headlines, INT8 quantization, rainflow cross-validation, dashboard
   drilldown, citation status, reproducibility).
4. The existing 附件 D 修訂說明 is renamed 附件 E and gets a new v2.2 row.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parent.parent
TEMPLATE = Path(r"C:\Users\user\Desktop\NYCU\ATCC\Sysblade_HyperBuffer_Proposal_v2.1.docx")
OUTPUTS = [
    REPO / "docs" / "proposal_v2.2_additions" / "Sysblade_HyperBuffer_Proposal_v2.2.docx",
    Path(r"C:\Users\user\Desktop\NYCU\ATCC\Sysblade_HyperBuffer_Proposal_v2.2.docx"),
]
TODAY = date.today().isoformat()


def _find_paragraph_by_text(doc, contains: str) -> int:
    """Return paragraph index whose text contains `contains`. -1 if none."""
    for i, p in enumerate(doc.paragraphs):
        if contains in p.text:
            return i
    return -1


def _apply_style_via_xml(paragraph_elem, style_id: str):
    """Set a paragraph's style by injecting <w:pStyle w:val="..."/> into pPr.

    Bypasses python-docx's styles[name] lookup which can be fragile when the
    style name and styleId differ (e.g. 'Heading 1' name vs 'Heading1' id).
    """
    pPr = paragraph_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph_elem.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


# Map human-readable style name → styleId from v2.1.docx scan
STYLE_ID = {
    "Heading 1": "Heading1",
    "Heading 2": "Heading2",
    "Heading 3": "Heading3",
    "List Paragraph": "ListParagraph",
}


def _insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    """Insert a new paragraph right AFTER `paragraph` and return it."""
    new_p_elem = OxmlElement("w:p")
    paragraph._element.addnext(new_p_elem)
    new_p = Paragraph(new_p_elem, paragraph._parent)
    if style:
        _apply_style_via_xml(new_p_elem, STYLE_ID.get(style, style))
    if text:
        new_p.add_run(text)
    return new_p


def _insert_paragraph_before(paragraph, text: str = "", style: str | None = None, bold: bool = False):
    """Insert a new paragraph right BEFORE `paragraph`."""
    new_p_elem = OxmlElement("w:p")
    paragraph._element.addprevious(new_p_elem)
    new_p = Paragraph(new_p_elem, paragraph._parent)
    if style:
        _apply_style_via_xml(new_p_elem, STYLE_ID.get(style, style))
    if text:
        run = new_p.add_run(text)
        run.bold = bold
    return new_p


def _set_run_bold(run, bold: bool = True):
    run.bold = bold
    return run


# ---------------------------------------------------------------------------
# Edit 1 — Cover: insert v2.2 修訂版 marker + URLs after the subtitle line.
# ---------------------------------------------------------------------------
def edit_cover(doc):
    # The subtitle "AI 時代的混合能量緩衝 BBU × 智能維運平台" sits at paragraph 3.
    # We insert a v2.2 block immediately after.
    idx = _find_paragraph_by_text(doc, "AI 時代的混合能量緩衝")
    if idx < 0:
        print("  WARN: cover subtitle not found — skipping cover edit")
        return
    anchor = doc.paragraphs[idx]

    # Add three lines after the subtitle:
    p3 = _insert_paragraph_after(anchor, f"v2.2 修訂版 ({TODAY}) — 軟體 demo 三件套已 live + 6 個關鍵數字 measured")
    p3.runs[0].bold = True
    p3.runs[0].font.color.rgb = RGBColor(0xC8, 0x16, 0x16)  # red accent

    p2 = _insert_paragraph_after(anchor, "GitHub: https://github.com/aericheng/atcc-sysblade")
    p1 = _insert_paragraph_after(anchor, "Live demo: https://sysblade-atcc.vercel.app")
    p1.runs[0].font.color.rgb = RGBColor(0x16, 0x4F, 0xC8)
    p2.runs[0].font.color.rgb = RGBColor(0x16, 0x4F, 0xC8)


# ---------------------------------------------------------------------------
# Edit 2 — Abstract: append a measured-numbers callout paragraph at end of A.
# ---------------------------------------------------------------------------
def edit_abstract(doc):
    # 找 A. 摘要的「一句話總結」末段,在它之後插 v2.2 measured callout
    idx = _find_paragraph_by_text(doc, "一句話總結")
    if idx < 0:
        # fallback: find the start of B.
        idx = _find_paragraph_by_text(doc, "B. 背景與目的") - 1
    if idx < 0:
        print("  WARN: abstract anchor not found — skipping")
        return
    # 插在 idx+1 之後(一句話總結的內文後),把新段落塞在 B. 之前
    insert_after = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else doc.paragraphs[idx]

    bullets = [
        "v2.2 measured 數字 — RUL 預測達 8.38 % MAPE(bagged-GBT,首次低於 v2.1 附件 B「< 10 %」承諾,Severson paper 9.1 % 對齊基線);LSTM fleet 部署 19.10 % MAPE 跨兩個 regime 誠實揭露",
        "INT8 量化實證 — LSTM ONNX 從 219 KB 壓到 63 KB(3.49× compression measured)、Δ MAPE 僅 +0.10 pp、R² 0.86 不變,屬 STM32N6 NPU 部署 go/no-go 證據",
        "獨立物理交叉驗證 — Rainflow + Wang 2011 第二條物理路徑於 worst-case GB200 工作點驗證 LIC 削峰使 LFP per-Ah 損傷下降 5.5 %(η = 0.945),不再單靠 Attia/Severson 統計外推",
        "軟體 demo 三件套已 live — /twin 數位孿生、/tco 計算器、/dashboard 機隊儀表板均部署於 Vercel,/dashboard 並支援 row click drilldown(SOH / RUL / 熱與操作層 metrics)",
        "Repo 公開展示 — github.com/aericheng/atcc-sysblade,CI 自動 cross-check 25 條數字 vs JSON ground truth,確保白皮書數字未被事後調整",
    ]
    # Insert in reverse so order ends up correct
    title = _insert_paragraph_after(insert_after, "v2.2 已交付 measured 重點(對應附件 D 完整實證):")
    title.runs[0].bold = True
    anchor = title
    for b in bullets:
        anchor = _insert_paragraph_after(anchor, "• " + b)


# ---------------------------------------------------------------------------
# Edit 3 — Insert new 附件 D 技術交付物實證 BEFORE existing 附件 D 修訂說明.
# ---------------------------------------------------------------------------
def edit_appendix(doc):
    idx = _find_paragraph_by_text(doc, "附件 D. v2.1 修訂說明")
    if idx < 0:
        print("  WARN: appendix D anchor not found — skipping appendix insert")
        return
    anchor = doc.paragraphs[idx]

    # Rename existing 附件 D 修訂說明 → 附件 E 修訂說明
    anchor.text = ""
    run = anchor.add_run("附件 E. 修訂說明 (v2.0 → v2.1 → v2.2)")
    # rename: text replaced; but Heading 1 style is preserved on the paragraph

    # Insert NEW 附件 D BEFORE the renamed 附件 E.
    _ins_before = _insert_paragraph_before

    # We build the appendix top-down — keep a moving anchor so each new
    # paragraph lands AFTER the previous one (which means we keep inserting
    # BEFORE the original 附件 E, which pushes new content above).
    title = _ins_before(anchor, "附件 D. v2.2 技術交付物實證", style="Heading 1")

    intro = _ins_before(anchor, "本附件記錄 v2.1 → v2.2 期間實際完成的技術交付物與 measured 數據,所有報告 JSON / 量化檔均存於公開 GitHub repo,可直接核驗。")

    # D.1 ML headlines
    _ins_before(anchor, "D.1 RUL 預測管線 measured 數字", style="Heading 2", bold=True)
    _ins_before(anchor, "Severson 2019 公開 124 顆 LFP 18650 cell;HDF5 解析路徑取出 138 顆有 ≥ 100 cycle 觀測。13-feature paper-aligned Full model 跨 5 種 regressor × 4 個 cell filter 全 sweep,結果存 data/processed/severson_model_eval.json:")
    bullets_d1 = [
        "Random split (10-seed median):bagged-GBT (K=24) + xstrict cell filter (cycle_life ≥ 400, n=134) MAPE 8.38 %、R² 0.890、per-seed [5.93, 12.91],7/10 seeds < 10 % — 首次達 v2.1 附件 B 軟體技術棧「< 10 % MAPE」承諾",
        "Cross-batch (b1+b2 → b3 新 protocol):bagged-OLS + xstrict MAPE 13.87 %、R² +0.207 — GBT 在 cross-batch 反退化(17–22 %)因 protocol-specific 過擬合,部署 routing 規則為「同 protocol 用 bagged-GBT、新 protocol fall back bagged-OLS」",
        "Cross-chemistry (Severson → NASA NMC):5/5 feature 全部 OOD、z-distance 5–65 σ — 量化證明跨化學體系需 per-chemistry calibration cycle,不可線性外插",
        "LSTM fleet 部署:Severson 138 + 50 顆 Severson-anchored synthetic BBU-duty cell(analytic decay + per-cell noise,非 PyBaMM 物理 aging)= 188 cells 訓練;test MAPE 19.10 % / R² 0.86 跨兩個 regime 誠實揭露 — 合成 cell 僅作 regime augmentation,production 信賴度仍以 Severson 真實 cells 為主;fleet 推論 production 已切 TCN(18.15 %,LSTM 為 baseline),學術 baseline 報 GBT",
    ]
    for b in bullets_d1:
        _ins_before(anchor, "• " + b)

    # D.2 INT8 quantization
    _ins_before(anchor, "D.2 INT8 動態量化實證(STM32N6 NPU 部署 go/no-go 證據)", style="Heading 2", bold=True)
    _ins_before(anchor, "scripts/quantize_lstm_onnx.py 用 onnxruntime.quantization.quantize_dynamic 真實量化,在 188-cell test 集量測:size 從 219.18 KiB 壓到 62.87 KiB(3.49× compression measured)、test MAPE 19.10 → 19.20 %(Δ MAPE +0.10 pp、R² 0.862 不變、平均預測偏移 0.57 %)、CPU INT8 對 FP32 1.11× 加速。STM32N6 Neural-ART NPU 1.6 MB FLASH 僅用 4 %。NPU 真機 latency 估算 27–109 µs(靜態 graph + ±2× 不確定區間),實機 trace SOP 寫於 docs/x_cube_ai_install_sop.md,需 ST 帳號 + X-CUBE-AI 9.x GUI 跑。")

    # D.3 Rainflow + Wang 2011 cross-validation
    _ins_before(anchor, "D.3 獨立物理交叉驗證 — Rainflow + Wang 2011", style="Heading 2", bold=True)
    _ins_before(anchor, "v2.1 §B.1 引述 Attia 2020 Nature [13] / Severson 2019 衰減模型外推之「~25 % LFP 循環壽命延長」屬統計外推,且主要組成是 BBU 低 duty 排程(§G.3 duty_factor 0.33,~50 cyc/yr vs Severson 1C/1C lab cadence),非 hybrid 拓樸貢獻。v2.2 補第二條獨立物理路徑量化 hybrid 拓樸貢獻:對 PyBaMM 產出的 LFP cell 電流跑 ASTM E1049-85 4-point rainflow + Wang 2011 J. Power Sources 半經驗 cycle-aging 公式,獨立估算每 Ah 損傷比 η = Q_loss(hybrid) / Q_loss(LFP-only):")
    bullets_d3 = [
        "demo 波形(±30 % / 100 ms):η = 1.012 — Wang kernel 在 0.5–6 C flat,demo cell C-rate 落 3.2–6 C 讓 hybrid 平直波形 per-Ah 損傷略高;主動揭露 demo 振幅本不是 hybrid 真正發揮優勢的工作點",
        "worst_case(10 C × 30 ms 脈衝,GB200 power-swing context per Choukse 2025 [11];團隊依 paper §IV-B per-cell BBU 下尺度推導):η = 0.945 — 5.5 % per-Ah 損傷下降,LIC 把 10 C 脈衝吸收後 LFP 看到的最大 C-rate 降至 4.8 C,對應 v2.1 §B.1「10–30 ms 5–10 C 瞬態」設計對象",
        "兩條獨立路徑同方向但量級不同:Attia/Severson 統計外推(BBU duty schedule 主導)貢獻大部分 ~25 % 壽命優勢;Wang/Rainflow 物理半經驗(hybrid 拓樸 per-Ah 損傷)額外貢獻 worst-case ~5.5 %、demo 波形近 neutral。兩條一起才完整支撐「~25 %」結論;報告 JSON 存 apps/web/public/scenarios/aging_rainflow_validation.json(repo 可追溯,不被 UI 消費)",
    ]
    for b in bullets_d3:
        _ins_before(anchor, "• " + b)

    # D.4 Dashboard drilldown
    _ins_before(anchor, "D.4 軟體三件套交付狀態", style="Heading 2", bold=True)
    _ins_before(anchor, "三件套均部署於 Vercel(github.com/aericheng/atcc-sysblade,main push 自動 build):")
    bullets_d4 = [
        "/twin Battery Digital Twin — PyBaMM DFN(Prada2013 LFP)+ LSTM RUL + 90 % MC-Dropout PI 經 split conformal 縮窄 44 %(raw 1910 → conformal 1075 cycles,test coverage 100 %);Inference Walkthrough 9 顆精選 cell span healthy / warning / early_aging / critical 四個狀態",
        "/tco TCO Calculator — 4 個 slider × 3 個 preset(Mid-tier · TX、Hyperscale · VA、Edge AI · PNW),即時計算 10 年 TCO saving;Hyperscale 500 racks · VA 年省 USD 482.9 k、payback 2.3 yr",
        "/dashboard Fleet Dashboard — 1000 台 seeded RNG 模擬,US fleet map + Tier-1/2/3 服務分層 + per-device drilldown(v2.2 新增,點擊 row 開 modal 顯示該 device SOH / RUL / 熱 / 操作層 metrics);全頁 SIMULATED DATA 浮水印 + persona disclaimer banner 雙重標註",
    ]
    for b in bullets_d4:
        _ins_before(anchor, "• " + b)

    # D.5 Reproducibility CI
    _ins_before(anchor, "D.5 Reproducibility CI gate", style="Heading 2", bold=True)
    _ins_before(anchor, "GitHub Actions .github/workflows/check.yml 在每次 push / PR 跑 4 個 job:apps/web typecheck + lint + Next.js build + scripts/check_whitepaper_numbers.py 數字 cross-check(25 條 assertion 掃白皮書 / README / battery-twin README vs JSON ground truth)。Whitepaper 數字若與 JSON 漂移,CI 紅燈擋 push,確保「8.38 %」「13.87 %」「19.10 %」「3.49×」「+0.10 pp」等所有 measured 數字隨時對齊原始量測,杜絕事後調整可能。")

    # D.6 Citation backing
    _ins_before(anchor, "D.6 Citation 實證進度", style="Heading 2", bold=True)
    _ins_before(anchor, "v2.1 引用之外部來源於 v2.2 期間逐條核實:")
    bullets_d6 = [
        "[v] Choukse 2025 (Microsoft + NVIDIA) arXiv:2508.14318 — 標題、作者、§III/§IV-B 內容定位均經 arxiv.org 全文確認(原 v2.1 引述為「NVIDIA reference」需修正為 Microsoft + NVIDIA collaboration)",
        "[v] Eaton XLR-48-166 module datasheet — 規格 48.6 V / 166 F / 54 Wh / ESR 5 mΩ 經 eaton.com 官方 datasheet 確認,正式 SKU XLR-48R6167-R",
        "[v] 三家競品 2024 年財報基準 — Eaton USD 24.9 B / Vertiv USD 8.012 B / Schneider EUR 38.15 B,均經公開 annual report 與 marketscreener 交叉確認",
        "(!) 仍待手動實證(複賽前 1 週逐條核實):Severson 2019 Figure 2c / Table 1 具體 panel 編號、Wang 2019 Prog. Energy Combust. Sci. §2.1 Table 2 LFP 230–270 °C / NMC 150–210 °C、Bandhauer 2011 §3、NASA PCoE B0005-B0018 cell rated 規格、Marquis 2019 §benchmark、ST AN5354 §Performance、BloombergNEF Lithium-Ion Battery Price Survey LFP 價格基準",
    ]
    for b in bullets_d6:
        _ins_before(anchor, "• " + b)


# ---------------------------------------------------------------------------
# Edit 4 — 修訂表(table 8): add v2.2 row
# ---------------------------------------------------------------------------
def edit_revisions_table(doc):
    # Find the 修訂 table — 8th table (index 8) per scan.
    # Rows: 5 修正 + 1 header. Add a new row for v2.2.
    if len(doc.tables) < 9:
        print("  WARN: revisions table not found — skipping")
        return
    tbl = doc.tables[8]
    new_row = tbl.add_row()
    cells = new_row.cells
    cells[0].text = "v2.2"
    cells[1].text = "v2.1 → v2.2"
    cells[2].text = "軟體 demo 三件套已 live + 6 個 measured 數字達標(8.38 % MAPE / +0.10 pp INT8 / η=0.945 rainflow / 3.49× compression / drilldown shipped / repo public + CI gate)"
    cells[3].text = f"見附件 D 完整實證({TODAY})"


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------
def main() -> int:
    if not TEMPLATE.exists():
        print(f"FATAL: template not found at {TEMPLATE}")
        return 1
    doc = Document(str(TEMPLATE))
    print(f"Loaded {TEMPLATE.name} — {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    print("Edit 1: cover...")
    edit_cover(doc)
    print("Edit 2: abstract callout...")
    edit_abstract(doc)
    print("Edit 3: 附件 D 新增...")
    edit_appendix(doc)
    print("Edit 4: 修訂表 add v2.2 row...")
    edit_revisions_table(doc)

    for out in OUTPUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        size_kb = out.stat().st_size / 1024
        print(f"  saved {out} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
